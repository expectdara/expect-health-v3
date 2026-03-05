#!/usr/bin/env python3
"""
Build utah-pt-reviewer-checklist.docx — Utah OAIP pilot version of the reviewer checklist.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

PURPLE = RGBColor(0x4C, 0x2C, 0x84)
DARK = RGBColor(0x1a, 0x1a, 0x2e)
GRAY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN_BG = 'E8F5E9'
AMBER_BG = 'FFF8E1'
RED_BG = 'FFEBEE'
PURPLE_HEX = '4C2C84'

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PURPLE
    return h

def add_para(text, bold=False, italic=False, size=None, color=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bold_prefix(bold_text, rest, size=11):
    p = doc.add_paragraph()
    r1 = p.add_run(bold_text)
    r1.bold = True
    r1.font.size = Pt(size)
    r2 = p.add_run(rest)
    r2.font.size = Pt(size)
    return p

def add_check_item(text, size=11):
    p = doc.add_paragraph()
    r = p.add_run('\u25A0 ')
    r.font.size = Pt(size)
    r2 = p.add_run(text)
    r2.font.size = Pt(size)
    return p

def add_callout(text, bg_color=GREEN_BG, bold_prefix=None):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{bg_color}"/>')
    pPr.append(shd)
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="360" w:right="360"/>')
    pPr.append(ind)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders_xml = f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>
    </w:tblBorders>'''
    tblPr.append(parse_xml(borders_xml))
    # Header row
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{PURPLE_HEX}"/>'))
        p = cell.paragraphs[0]; p.clear()
        run = p.add_run(h); run.bold = True; run.font.color.rgb = WHITE; run.font.size = Pt(10)
    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            p = cell.paragraphs[0]; p.clear()
            run = p.add_run(cell_text); run.font.size = Pt(10)
    return table


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ══════════════════════════════════════════════════════════════════════════════

# Header
p = doc.add_paragraph()
run = p.add_run('CLINICAL REVIEWER CHECKLIST')
run.bold = True; run.font.size = Pt(10); run.font.color.rgb = PURPLE
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
run = p.add_run('PT Review Guide — Utah OAIP Pilot')
run.bold = True; run.font.size = Pt(24); run.font.color.rgb = DARK
p.paragraph_format.space_after = Pt(4)

add_para('Dr. Nicole Dugan, PT, DPT, WCS | Expect Clinical Platform — Utah OAIP Regulatory Sandbox', size=11, color=GRAY, space_after=8)

add_callout(
    'The system handles scoring, ICD-10 diagnosis, and exercise selection automatically. Your role is to verify logic matched correctly, review OAIP safety flags, confirm provider routing, and approve or modify the plan before it reaches the patient. Keep reviews under 5 minutes. A review timer is displayed at the top of each case.',
    bg_color='F3E8FF'
)

# ── Step 1 ──
add_heading('Step 1: Safety Verification', level=2)
add_check_item("Confirm no 'Yes' answers for red flags (bleeding, fever, chest pain, headache with vision changes, UTI symptoms). System should have redirected -- quick visual check for glitches.")
add_check_item("If a SAFETY_ANSWER_CHANGED alert is present: review the change details. Verify the patient's revised answer is clinically appropriate. The original and revised answers are logged for OAIP audit.")

# ── Step 2 ──
add_heading('Step 2: Tier Assignment', level=2)
add_check_item('Check ICIQ Total and verify correct tier:')
doc.add_paragraph('')
add_table(
    ['ICIQ', 'Tier', 'Program'],
    [
        ['13-21', 'Beginner', '12 wks, starts with Supine PF Activation, PT review q3 days'],
        ['6-12', 'Moderate', '8 wks, standard Quick Squeeze start, weekly check-in'],
        ['1-5', 'Advanced', '6 wks, higher volume, functional positions, bi-weekly check-in'],
    ]
)
doc.add_paragraph('')

# ── Step 3 ──
add_heading('Step 3: Adjunct Verification', level=2)
add_table(
    ['Trigger', 'Should Include', '\u25A0'],
    [
        ['Pain Composite >= 4', 'Pain De-Sensitization Breathing as first exercise', '\u25A0'],
        ['Pain >= 5 OR Dyspareunia', 'Vaginal Dilator Therapy adjunct', '\u25A0'],
        ['Urgency or Mixed subtype', 'Bladder Retraining adjunct', '\u25A0'],
        ['Constipation >= 2', 'Bowel Management Program', '\u25A0'],
        ['Med modification = Yes', 'Medication Review Referral adjunct (DIAPPERS framework)', '\u25A0'],
    ]
)
doc.add_paragraph('')

# ── Step 4 ──
add_heading('Step 4: Cue Personalization Check', level=2)
add_check_item("Confirm exercise instructions use the patient's selected cue style (biologic / imaginative / breathing / simple contract / default).")
add_check_item("Instructions should NOT say only 'squeeze your pelvic floor' or 'engage your pelvic floor' -- research shows these general cues are least effective (Crane & Dugan, 2025).")
add_check_item("If biologic cue selected: verify the urination warning appears (\"Never practice this during actual urination\").")

# ── Step 5 ──
add_heading('Step 5: Avoidance and Escalation Check', level=2)
add_check_item('If >= 3 avoided activities: verify HIGH IMPACT flag is set and goal includes resuming avoided activities.')
add_check_item('If ICIQ >= 19 OR (>= 3 avoided activities + elevated pain): verify enhanced monitoring is flagged.')
add_check_item('If medication modification = Yes: verify Medication Review Referral adjunct is present.')
add_check_item('If pudendal neuralgia indicators present (sitting_long trigger + pain >6/10): verify ICD-10 G57.91 diagnosis and "avoid prolonged sitting exercises" precaution are present.')

# ── Step 5b ──
add_heading('Step 5b: Prenatal Protocol Check (if applicable)', level=2)
add_check_item('If patient indicated active pregnancy: verify PRENATAL_FLAG is set and prenatal exercise substitutions have been applied.')
add_check_item("Confirm incline/side-lying variants are appropriate for the patient's current trimester.")
add_check_item('Verify the vena cava compression precaution block is present.')
add_check_item('Confirm the care plan is labeled "Prenatal Pelvic Floor Protocol."')

# ── Step 5c ──
add_heading('Step 5c: Depression Screening Review', level=2)
add_check_item('If PHQ-2 score >= 3/6: verify depression_screen_positive flag is present (MODERATE 3-4, HIGH 5-6).')
add_check_item("Review the inline clinical consistency warning. Verify patient's capacity to consent is appropriate for continued care.")
add_check_item('Flag is surfaced as DEPRESSION_RISK in the OAIP Red Flags dashboard.')

# ── Step 6 ──
add_heading('Step 6: Provider Routing & Fax Verification', level=2)
add_check_item('Verify the provider fax verification status:')
doc.add_paragraph('')
add_table(
    ['Status', 'Badge', 'Action Required'],
    [
        ['NPI-Verified', 'Green VERIFIED', 'None -- auto-populated from CMS NPI Registry.'],
        ['Default', 'Amber DEFAULT', 'Confirm the pre-populated fax is correct for the provider\'s organization and specialty. Review the matching rule.'],
        ['Manual Entry', 'Amber MANUAL ENTRY', 'Verify fax number before sending. Not validated against NPI Registry.'],
        ['Concierge Pending', 'Hold Banner', 'Transmission is on hold. Do NOT send fax until provider is verified.'],
    ]
)
doc.add_paragraph('')
add_check_item('If provider was submitted via concierge flow: confirm "HOLD -- Pending Destination Verification" banner is active and Send Fax button is disabled.')

# ── Step 7 ──
add_heading('Step 7: Encounter Note Review', level=2)
add_check_item('Review the auto-generated encounter note (CMS-compliant SOAP format). Verify accuracy.')
add_check_item('Use shorthand expansion (e.g., "PFM" auto-expands to "pelvic floor muscle") for efficient editing.')
add_check_item('Verify clinical alerts are prepended to the note (medication modification, safety answer changes, depression screening, prenatal protocol).')

# ── Step 8 ──
add_heading('Step 8: Final Sign-Off', level=2)
add_check_item("Does the overall plan make clinical sense for this patient's presentation?")
add_check_item('Any red flags the system may have missed?')
add_check_item('Are exercise parameters (sets, reps, hold times) appropriate for the assigned tier?')

doc.add_paragraph('')
add_callout("If everything checks out: Click 'Approve Plan.' Target review time: 3-5 minutes.", bg_color=GREEN_BG, bold_prefix='If everything checks out: ')
add_callout("If something needs adjustment: Modify the plan directly in the review interface. All modifications are logged in the audit trail.", bg_color=AMBER_BG, bold_prefix='If something needs adjustment: ')

# ── HIPAA Reminder ──
doc.add_paragraph('')
add_heading('HIPAA PHI Handling Reminders', level=2)
add_callout(
    'All patient data visible in the review interface is Protected Health Information (PHI). '
    'Do not copy PHI to external systems, emails, or personal notes. '
    'Use the built-in Auditor Mode toggle when sharing screens or demonstrating the platform. '
    'The audit log records every action you take, including review timestamps and modifications. '
    'If you need to discuss a case externally, use de-identified references only.',
    bg_color=RED_BG
)

# Footer
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Utah OAIP Regulatory Sandbox Pilot | CONFIDENTIAL')
run.font.size = Pt(9); run.font.color.rgb = GRAY; run.italic = True
p2 = doc.add_paragraph()
run2 = p2.add_run('This checklist corresponds to the prevention-first safety architecture described in Appendix B of the OAIP pilot proposal and the OAIP Red Flags dashboard described in Section 1.4.')
run2.font.size = Pt(9); run2.font.color.rgb = GRAY; run2.italic = True

# Save
output_path = '/Users/Dara Cook/Downloads/expect-utah/utah-pt-reviewer-checklist.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
