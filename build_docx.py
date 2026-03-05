#!/usr/bin/env python3
"""
Convert expect_oaip_pilot_proposal_REWRITTEN.txt into a properly formatted .docx
matching the original document's styles, table formatting, and layout.
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Read the rewritten text ──────────────────────────────────────────────────
with open('/Users/Dara Cook/Downloads/expect-utah/expect_oaip_pilot_proposal_REWRITTEN.txt', 'r') as f:
    full_text = f.read()

# ── Create document from original as template (to inherit styles) ────────────
doc = Document('/Users/Dara Cook/Downloads/expect-utah/expect_oaip_pilot_proposal_updated (1).docx')

# Clear all existing content but preserve sectPr
body = doc.element.body
sectPr = body.find(qn('w:sectPr'))
for child in list(body):
    if child.tag != qn('w:sectPr'):
        body.remove(child)
if sectPr is None:
    sectPr = parse_xml(
        f'<w:sectPr {nsdecls("w")}>'
        '<w:pgSz w:w="12240" w:h="15840"/>'
        '<w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440" '
        'w:header="720" w:footer="720" w:gutter="0"/>'
        '</w:sectPr>'
    )
    body.append(sectPr)

PURPLE = '4C2C84'

# ── Helper functions ─────────────────────────────────────────────────────────

def add_paragraph(text, style='Normal', alignment=None, bold=None, space_after=None):
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    if bold is not None:
        run = p.add_run(text)
        run.bold = bold
    else:
        p.add_run(text)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bold_prefix_paragraph(bold_part, rest, style='Normal'):
    p = doc.add_paragraph(style=style)
    run_bold = p.add_run(bold_part)
    run_bold.bold = True
    if rest:
        p.add_run(rest)
    return p

def add_table(headers, rows):
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders_xml = f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>
    </w:tblBorders>'''
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    tblPr.append(parse_xml(borders_xml))
    header_row = table.rows[0]
    for ci, header_text in enumerate(headers):
        cell = header_row.cells[ci]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{PURPLE}"/>')
        existing_shd = tcPr.find(qn('w:shd'))
        if existing_shd is not None:
            tcPr.remove(existing_shd)
        tcPr.append(shading)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(header_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(cell_text)
            run.font.size = Pt(10)
    return table

def parse_markdown_table(lines):
    headers = []
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith('|'):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if all(set(c) <= {'-', ' '} for c in cells):
            continue
        if not headers:
            headers = cells
        else:
            rows.append(cells)
    return headers, rows

# ── Parse tables from the "Tables" section at the end ────────────────────────
lines = full_text.split('\n')

table_section_start = None
for idx, line in enumerate(lines):
    if line.strip() == 'Tables':
        table_section_start = idx
        break

tables_data = {}
if table_section_start is not None:
    tidx = table_section_start + 1
    current_table_name = None
    current_table_lines = []
    while tidx < len(lines):
        line = lines[tidx].strip()
        m = re.match(r'^Table (\d+):', line)
        if m:
            if current_table_name is not None and current_table_lines:
                h, r = parse_markdown_table(current_table_lines)
                if h:
                    tables_data[current_table_name] = (h, r)
            current_table_name = int(m.group(1))
            current_table_lines = []
        elif line.startswith('|'):
            current_table_lines.append(line)
        tidx += 1
    if current_table_name is not None and current_table_lines:
        h, r = parse_markdown_table(current_table_lines)
        if h:
            tables_data[current_table_name] = (h, r)

tables_inserted = set()
end_line = table_section_start if table_section_start else len(lines)

# ── Known appendix section titles (EXPECT Heading 1) ─────────────────────────
appendix_section_titles = [
    '0. Intervention Risk Profile',
    '1. Prevention-First Safety Architecture',
    '2. Incident Reporting and Expedited Clinical Review',
    '3. Root-Cause Classification',
    '4. Corrective Care Commitment',
    '5. Pilot Program Terms (Defined and Bounded)',
    '6. Reporting and Continuous Improvement',
    '1. Overview',
    '2. Data Collected',
    '3. Data Storage and Security',
    '4. De-Identification Standard',
    '5. Data Sharing: Utah OAIP Oversight',
    '6. Data Sharing: IRB-Approved Research',
    '7. Data Sharing: External Auditors',
    '8. Data Sharing: Anonymized Public Dashboards',
    '9. Participant Rights',
    '10. Data Governance Contact',
    '1. The Access Gap in Utah',
    '2. Clinical Evidence Base',
    '3. Pilot Phases and PT Review Evolution',
    '4. Success and Failure Definitions',
    '5. End-to-End Patient Experience (HIPAA-Compliant)',
    '6. Vulnerable Populations Considerations',
    '7. Platform Quality Improvements (February -- March 2026)',
]

h1_patterns = [
    'Executive Summary',
    '1. About You',
    '2. Regulatory Relief Request',
    '3. Pilot Design',
    '4. Risk Management',
    '5. Feasibility',
    '6. Public Relations',
    '7. Confidentiality Claims',
]

bold_headings = [
    'Patient Journey', 'Patient Journey Overview', 'AI Disclosure at Every Touchpoint',
    'PT Provider Experience', 'OAIP Oversight Experience', 'Right to Opt Out',
    'Data Collected and Stored', 'Data Usage', 'EHR Interface', 'HIPAA Compliance',
    'Expertise', 'Funding', 'Liability Insurance', 'Stakeholder Networks',
    'Academic Literature', 'Pre-Deployment Testing', 'Related Service/Product',
    'Federal Research Partnership', 'Proposed Reporting Process',
    'Trade Secrets', 'Commercial Information', 'Reasons Supporting Confidentiality',
    'How Exclusions Are Screened', 'Researchers', 'External Auditors',
    'Lower Risk Profile Than Approved Precedent',
    'Prenatal Pelvic Floor Protocol',
    'Leadership Team',
]

big_picture_prefixes = ['Problem Statement:', 'Prevalence:', 'Clinician Scarcity:',
                        'Harms of Inaction:', 'Low-Risk Intervention:', 'Proposed Solution:',
                        'Billing Strategy:', 'Codes Utilized:', 'Payer Strategy:', 'HSA/FSA Eligibility:',
                        'Enrollment Projections:']

# ── Render document ──────────────────────────────────────────────────────────
i = 0
while i < end_line:
    line = lines[i]
    stripped = line.strip()

    if stripped == '':
        i += 1
        continue

    # Title block
    if stripped == 'EXPECT' and i < 5:
        add_paragraph('EXPECT ', style='EXPECT Title', alignment=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1; continue
    if stripped == 'AI-Augmented Pelvic Floor Physical Therapy Platform' and i < 10:
        add_paragraph(stripped, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1; continue
    if stripped == '------' or stripped.startswith('───'):
        add_paragraph('──────────────────────────────', alignment=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1; continue
    if stripped == 'Pilot Proposal to the Utah Office of Artificial Intelligence Policy':
        add_paragraph(stripped, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1; continue
    if stripped == 'Regulatory Mitigation Agreement Application':
        add_paragraph(stripped, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1; continue
    if stripped == 'Under the Utah Artificial Intelligence Policy Act (UAIPA)':
        add_paragraph(stripped, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1; continue
    if stripped == 'February 2026':
        add_paragraph('')
        add_paragraph(stripped, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1; continue
    if stripped.startswith('CONFIDENTIAL'):
        add_paragraph(stripped, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        i += 1; continue

    # EXPECT Heading 1 - main sections
    if stripped in h1_patterns:
        add_paragraph(stripped, style='EXPECT Heading 1')
        i += 1; continue

    # APPENDIX headings
    if stripped.startswith('APPENDIX ') and ':' in stripped:
        if stripped.startswith('APPENDIX A:') or stripped.startswith('APPENDIX D'):
            add_paragraph(stripped, style='EXPECT Heading 1')
        else:
            add_paragraph(stripped, style='EXPECT Title')
        i += 1; continue

    # EXPECT Heading 2 - numbered sub-sections (e.g., "1.1 What Is...")
    h2_pattern = re.match(r'^(\d+\.\d+)\s+(.+)$', stripped)
    if h2_pattern:
        add_paragraph(stripped, style='EXPECT Heading 2')
        i += 1; continue

    # Appendix sub-section titles → EXPECT Heading 1
    if stripped in appendix_section_titles:
        add_paragraph(stripped, style='EXPECT Heading 1')
        i += 1; continue

    # Section 3.5 sub-headings (FDA, OAIP, Public)
    if stripped.startswith('FDA: '):
        add_paragraph('FDA', style='EXPECT Title')
        add_paragraph(stripped[5:])
        i += 1; continue
    if stripped.startswith('OAIP: '):
        add_paragraph('OAIP', style='EXPECT Title')
        add_paragraph(stripped[6:])
        i += 1; continue
    if stripped.startswith('Public: '):
        p = doc.add_paragraph(style='Normal')
        run = p.add_run('Public')
        run.bold = True
        add_paragraph(stripped[8:])
        i += 1; continue

    # Bold sub-headings
    step_match = re.match(r'^Step \d+:', stripped)
    sub_sub = re.match(r'^\d+[a-z]\.\s', stripped)
    draft_heading = re.match(r'^\d+[a-z]?\.\s.+\(Draft\)$', stripped)
    if stripped in bold_headings or step_match or sub_sub or draft_heading:
        p = doc.add_paragraph(style='Normal')
        run = p.add_run(stripped)
        run.bold = True
        i += 1; continue

    # Appendix sub-titles
    if stripped.startswith('Expect -- Utah OAIP Pilot'):
        p = doc.add_paragraph(style='Normal')
        run = p.add_run(stripped)
        run.italic = True
        i += 1; continue

    # Table insertion
    see_table = re.match(r'^\(See Table (\d+) below', stripped)
    if see_table:
        tnum = int(see_table.group(1))
        add_paragraph(stripped)
        if tnum in tables_data and tnum not in tables_inserted:
            doc.add_paragraph('')
            h, r = tables_data[tnum]
            add_table(h, r)
            tables_inserted.add(tnum)
            doc.add_paragraph('')
        i += 1; continue

    # Bullet points
    if stripped.startswith('- '):
        bullet_text = stripped[2:]
        colon_match = re.match(r'^([^:]+):\s*(.+)$', bullet_text)
        if colon_match and len(colon_match.group(1)) < 60:
            p = doc.add_paragraph(style='Normal')
            run_b = p.add_run(colon_match.group(1) + ': ')
            run_b.bold = True
            p.add_run(colon_match.group(2))
        else:
            p = doc.add_paragraph(style='Normal')
            p.add_run(bullet_text)
        pPr = p._p.get_or_add_pPr()
        ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="720" w:hanging="360"/>')
        existing_ind = pPr.find(qn('w:ind'))
        if existing_ind is not None:
            pPr.remove(existing_ind)
        pPr.append(ind)
        if p.runs:
            p.runs[0].text = '\u2022\t' + p.runs[0].text
        i += 1; continue

    # Indented sub-items
    if line.startswith('  ') and stripped.startswith('- '):
        p = doc.add_paragraph(style='Normal')
        p.add_run('\u2022\t' + stripped[2:])
        pPr = p._p.get_or_add_pPr()
        ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="1080" w:hanging="360"/>')
        existing_ind = pPr.find(qn('w:ind'))
        if existing_ind is not None:
            pPr.remove(existing_ind)
        pPr.append(ind)
        i += 1; continue

    # Bold prefix patterns (Big Picture items)
    matched_bp = False
    for prefix in big_picture_prefixes:
        if stripped.startswith(prefix):
            add_bold_prefix_paragraph(prefix + ' ', stripped[len(prefix):].strip())
            matched_bp = True
            break
    if matched_bp:
        i += 1; continue

    # Numbered items
    num_item = re.match(r'^(\d+)\.\s+(.+)$', stripped)
    if num_item:
        p = doc.add_paragraph(style='Normal')
        p.add_run(stripped)
        pPr = p._p.get_or_add_pPr()
        ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="720" w:hanging="360"/>')
        existing_ind = pPr.find(qn('w:ind'))
        if existing_ind is not None:
            pPr.remove(existing_ind)
        pPr.append(ind)
        i += 1; continue

    # Email template content
    if stripped.startswith('Subject:') or stripped.startswith('Hi ['):
        p = doc.add_paragraph(style='Normal')
        run = p.add_run(stripped)
        if stripped.startswith('Subject:'):
            run.bold = True
        i += 1; continue
    if stripped.startswith('[Button:'):
        p = doc.add_paragraph(style='Normal')
        run = p.add_run(stripped)
        run.bold = True
        i += 1; continue

    # Default: Normal paragraph
    add_paragraph(stripped)
    i += 1

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = '/Users/Dara Cook/Downloads/expect-utah/expect_oaip_pilot_proposal_FINAL.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
