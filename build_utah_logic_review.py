#!/usr/bin/env python3
"""
Build utah-pt-logic-review.docx — Utah OAIP pilot version of the wellness logic review.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

PURPLE = RGBColor(0x4C, 0x2C, 0x84)
DARK = RGBColor(0x1a, 0x1a, 0x2e)
GRAY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
PURPLE_HEX = '4C2C84'

def add_h1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs: r.font.color.rgb = PURPLE
    return h

def add_h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs: r.font.color.rgb = PURPLE
    return h

def add_h3(text):
    h = doc.add_heading(text, level=3)
    for r in h.runs: r.font.color.rgb = PURPLE
    return h

def add_p(text, bold=False, italic=False, size=11, color=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold; run.italic = italic; run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bp(bold_text, rest, size=11):
    p = doc.add_paragraph()
    r1 = p.add_run(bold_text); r1.bold = True; r1.font.size = Pt(size)
    r2 = p.add_run(rest); r2.font.size = Pt(size)
    return p

def add_callout(text, bg='F3E8FF', bold_prefix=None):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{bg}"/>'))
    pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="360" w:right="360"/>'))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix:
        r1 = p.add_run(bold_prefix); r1.bold = True; r1.font.size = Pt(10)
        r2 = p.add_run(text); r2.font.size = Pt(10)
    else:
        r = p.add_run(text); r.font.size = Pt(10)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    borders = f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>
    </w:tblBorders>'''
    tbl.tblPr.append(parse_xml(borders))
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{PURPLE_HEX}"/>'))
        p = cell.paragraphs[0]; p.clear()
        run = p.add_run(h); run.bold = True; run.font.color.rgb = WHITE; run.font.size = Pt(10)
    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            p = cell.paragraphs[0]; p.clear()
            run = p.add_run(cell_text); run.font.size = Pt(10)
    return table


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ══════════════════════════════════════════════════════════════════════════════

# Title
p = doc.add_paragraph()
run = p.add_run('CLINICAL LOGIC REVIEW')
run.bold = True; run.font.size = Pt(10); run.font.color.rgb = PURPLE
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
run = p.add_run('Expect Utah OAIP Clinical Platform')
run.bold = True; run.font.size = Pt(24); run.font.color.rgb = DARK
p.paragraph_format.space_after = Pt(2)

add_p('Questionnaire Structure, Scoring Algorithms, Diagnosis Generation, and Exercise Recommendation Logic', size=12, color=GRAY, space_after=4)
add_p('Prepared for Dr. Nicole Dugan, PT, DPT, WCS | Utah OAIP Regulatory Sandbox Pilot | February 2026 | CONFIDENTIAL', size=10, color=GRAY, space_after=8)

add_callout(
    'Scoring tiers per Klovning et al. (2009) \u2014 ICIQ-UI SF severity classification validated against '
    'the Incontinence Severity Index (n=1,812). Exercise cueing methodology informed by Crane, Dugan et al. (2025) \u2014 '
    'the #9 most downloaded paper in the history of the Journal of Women\'s & Pelvic Health Physical Therapy. '
    'Dr. Dugan is co-author on this landmark research and Lead Clinical Expert for the Expect platform.',
    bold_prefix='Clinical Evidence Base: '
)

# ══════════════════════════════════════════════════════════════════════════════
# Section 1: Safety Screening
# ══════════════════════════════════════════════════════════════════════════════
add_h1('Section 1: Safety Red-Flag Screening (5 questions)')
add_p("Binary yes/no. Any 'Yes' immediately stops intake. No plan is generated. Patient is redirected to emergency or physician care.", space_after=6)

for item in [
    ("rf_bleed: ", "Unexplained vaginal bleeding? \u2192 Redirect to physician"),
    ("rf_fever: ", "Fever above 100.4\u00b0F? \u2192 Redirect to ER"),
    ("rf_chest: ", "Chest pain/difficulty breathing? \u2192 Redirect to ER"),
    ("rf_head: ", "Severe headache + vision changes? \u2192 Redirect to ER"),
    ("rf_uti: ", "Burning urination + blood + fever? \u2192 Redirect to physician"),
]:
    add_bp(item[0], item[1])

add_callout(
    'Safety gate: If a patient changes a previously positive safety answer, a confirmation modal is displayed. '
    'The change is logged as a SAFETY_ANSWER_CHANGED audit event and surfaced as a red flag in the PT review '
    'dashboard and OAIP Red Flags dashboard. The reviewing PT must verify the revised answer is clinically appropriate.',
    bg='FFEBEE'
)

# ══════════════════════════════════════════════════════════════════════════════
# Section 2: Exclusion Screening
# ══════════════════════════════════════════════════════════════════════════════
add_h1('Section 2: Eligibility / Exclusion Screening (6 + 1 conditional)')
add_p("Six unconditional exclusion questions are always shown. One conditional question (high-risk pregnancy) is shown only to patients who indicate active pregnancy on the demographics screen. Any 'Yes' on an unconditional question halts intake and redirects to physician/specialist. No AI plan is generated.", space_after=6)

exclusions = [
    ("ex_neuro: ", "Neurogenic bladder or known neurological conditions (MS, spinal cord injury, Parkinson's, stroke)? \u2192 Redirect to specialist"),
    ("ex_mesh: ", "Post-surgical mesh complications or post-surgical pelvic condition? \u2192 Redirect to specialist"),
    ("ex_prolapse: ", "Grade 3 or higher pelvic organ prolapse? \u2192 Redirect to specialist"),
    ("ex_malig: ", "Known or suspected pelvic malignancy (bladder, uterine, cervical, ovarian)? \u2192 Redirect to physician"),
    ("ex_infect: ", "Active pelvic infection, fistula, or abscess? \u2192 Redirect to physician"),
    ("ex_ic: ", "Interstitial cystitis with confirmed Hunner lesions? \u2192 Redirect to urologist"),
]
for item in exclusions:
    add_bp(item[0], item[1])

doc.add_paragraph('')
add_bp("ex_highrisk_preg (CONDITIONAL): ", '"Have you been told by your OB/GYN or midwife that you have a high-risk pregnancy or been advised to avoid exercise?" Shown only when prenatal_flag is true. Yes \u2192 Block with message: "Please consult your prenatal care team before starting any pelvic floor program. We\'re here when you get the green light." No \u2192 Continue with prenatal-adapted care plan.')

# ══════════════════════════════════════════════════════════════════════════════
# Section 3: Depression Screening
# ══════════════════════════════════════════════════════════════════════════════
add_h1('Section 3: Depression Screening \u2014 PHQ-2 (2 questions)')
add_p('Two validated questions screening for depression. Scoring: 0-6 scale, threshold >= 3.', space_after=6)
add_bp("phq2_q1: ", '"Over the past 2 weeks, how often have you been bothered by little interest or pleasure in doing things?" Not at all (0), Several days (1), More than half the days (2), Nearly every day (3)')
add_bp("phq2_q2: ", '"Over the past 2 weeks, how often have you been bothered by feeling down, depressed, or hopeless?" Same scale.')
doc.add_paragraph('')
add_bp("Scoring: ", "PHQ-2 Total = phq2_q1 + phq2_q2 (range 0-6)")
add_bp("Threshold: ", ">= 3 triggers depression_screen_positive audit event")
add_bp("Severity: ", "Score 3-4 = MODERATE, Score 5-6 = HIGH")
add_bp("Clinical action: ", "Inline consistency warning displayed. Case flagged for PT to verify capacity to consent. Surfaced as DEPRESSION_RISK in OAIP Red Flags dashboard.")

# ══════════════════════════════════════════════════════════════════════════════
# Section 4: Demographics & Provider Selection
# ══════════════════════════════════════════════════════════════════════════════
add_h1('Section 4: Demographics & Provider Selection')

add_h2('Demographics')
add_p('Patient provides: first name, last name, date of birth (validated 18-120), phone (auto-formatted), pregnancy status, delivery history.', space_after=6)
add_callout(
    'If pregnancy_status = "Currently pregnant": PRENATAL_FLAG is set on the patient record. '
    'A gentle inline note appears: "We\'ll tailor your care plan with prenatal-safe modifications so you can safely '
    'support your pelvic floor throughout your pregnancy." This is informational, not a blocker. '
    'A PRENATAL_PROTOCOL_APPLIED audit event is logged.',
    bg='E8F5E9'
)

add_h2('Provider Selection via NPI Registry')
add_p('Three pathways for provider selection:', space_after=4)
add_bp("1. NPI Registry lookup: ", "Patient searches by provider name. Results verified against CMS NPI Registry API (NPI-1, Utah filter). Typo-tolerant matching. Auto-populates provider credentials, specialty, fax number, and organization.")
add_bp("2. Demo/sandbox providers: ", "Pre-loaded test providers for platform demonstrations. Demo fax numbers are used as-is; production default fax rules are never applied.")
add_bp("3. Concierge verification: ", 'Patient manually enters provider details for providers not in the NPI Registry. Provider is held in "Pending Verification" status. No fax is transmitted until the provider is verified.')

add_h2('Fax Verification (3-Tier System)')
add_table(
    ['Tier', 'Source', 'Badge', 'Behavior'],
    [
        ['NPI-Verified', 'CMS NPI Registry API', 'Green VERIFIED', 'Auto-populated. No PT action required.'],
        ['Default', 'Org/specialty matching rules', 'Amber DEFAULT', 'Pre-populated when NPI has no fax on file. PT must verify before sending. DEFAULT_FAX_APPLIED audit event logged.'],
        ['Manual Entry', 'PT enters manually', 'Amber MANUAL ENTRY', 'PT types fax number. Not validated against NPI Registry. Warning displayed.'],
    ]
)
doc.add_paragraph('')
add_callout(
    'Concierge pending: If the provider was submitted via concierge flow and verification is still pending, '
    'fax transmission is held. The PT sees a "HOLD \u2014 Pending Destination Verification" banner, and the '
    'Send Fax button is disabled until the provider is verified.',
    bg='FFF8E1'
)

# ══════════════════════════════════════════════════════════════════════════════
# Section 5: Bladder & Leaking
# ══════════════════════════════════════════════════════════════════════════════
add_h1('Section 5: Bladder & Leaking Patterns (8 questions)')
add_p('Based on ICIQ-UI SF and ICIQ-FLUTS. Severity bands per Klovning et al. (2009).', space_after=6)

for item in [
    ("iciq1: ", "How often do you experience bladder leaking? Never(0) to Always(5)"),
    ("iciq2: ", "When you leak, how much? None(0), Small(2), Moderate(4), Large(6)"),
    ("iciq3: ", "How much does leaking affect everyday life? Slider 0-10"),
    ("iciq4: ", "When does leaking happen? (multi-select) urgency, stress_cough, stress_exercise, post_void, nocturnal, unknown, continuous"),
    ("fl2a: ", "Night bathroom trips? None(0) to 4+(4). >= 2 = Nighttime Waking flag"),
    ("fl3a: ", "Sudden urgent need to rush? Never(0) to All the time(4)"),
    ("fl5a: ", "Daytime bathroom frequency? 1-6x(0) to 13+(4)"),
    ("avoid_activities: ", "Are you avoiding activities because of pelvic symptoms? (multi-select) Exercise, Social events, Travel, Sexual activity, Lifting, None. >= 3 categories = HIGH IMPACT flag \u2192 enhanced monitoring"),
]:
    add_bp(item[0], item[1])

add_h2('Scoring: Bladder')
add_bp("ICIQ Total (it) ", "= iciq1 + iciq2 + iciq3 (range 0-21)")
add_bp("Daily Impact ", "= iciq3 directly")
add_p('3-Tier System (per Klovning et al. 2009):', bold=True, space_after=4)
add_table(
    ['ICIQ Range', 'Severity', 'Tier', 'Program'],
    [
        ['1-5', 'Slight', 'Advanced', '6 weeks, higher volume, functional exercises'],
        ['6-12', 'Moderate', 'Moderate', '8 weeks, standard protocol'],
        ['13-18', 'Severe', 'Beginner', '12 weeks, supine/gravity-neutral start'],
        ['19-21', 'Very Severe', 'Beginner', '12 weeks + enhanced monitoring flagged'],
    ]
)
doc.add_paragraph('')

add_bp("Incontinence Subtype Classification: ", "urgency AND stress \u2192 Mixed; urgency only \u2192 Urge; stress only \u2192 Stress; continuous \u2192 Continuous; else \u2192 General")
add_bp("Avoidance Impact: ", ">= 3 avoided activity categories \u2192 HIGH IMPACT flag \u2192 enhanced monitoring regardless of ICIQ score")

# ══════════════════════════════════════════════════════════════════════════════
# Section 6: Pain & Sensitivity
# ══════════════════════════════════════════════════════════════════════════════
add_h1('Section 6: Pain & Sensitivity (GUPI-Female + Pain Assessment)')
add_p('Based on GUPI-Female. Expanded from original with ICIQ-FLUTSsex integration.', space_after=6)

add_h2('GUPI-Female (24 questions)')
add_bp("gupi1a-1d: ", "Pain/discomfort locations: perineum, vaginal entrance, suprapubic, during/after sexual activity (4 yes/no)")
add_bp("gupi2a-2b: ", "Urinary symptoms: frequency, urgency")
add_bp("gupi3-4: ", "Frequency and severity of discomfort (0-5 scales)")
add_bp("gupi_qol1-3: ", "Quality of life impact subscales")
add_bp("Scoring: ", "GUPI Pain (P) 0-23, Urinary (U) 0-10, Quality (Q) 0-12. Total 0-45.")

add_h2('Pain Assessment')
add_bp("pain1: ", "Current pain level (0-10)")
add_bp("pain2: ", "Average pain level over past week (0-10)")
add_bp("pain3: ", "Effect on daily activities (0-10)")
add_bp("pain4: ", "Pain locations (multi-select, includes sitting, sexual activity)")
add_bp("Scoring: ", "Pain Composite = (pain1 + pain2) / 2 (range 0-10). Functional Impact = pain3.")

add_h2('Pudendal Neuralgia Detection')
add_callout(
    'If symptom_triggers includes "sitting_long" AND pain composite > 6/10: '
    'System flags suspected pudendal neuralgia. Triggers ICD-10 code G57.91, '
    'adds precaution against prolonged sitting exercises, and activates a mandatory '
    'PT review guardrail alert. The PT evaluates nerve involvement before progressing the exercise program.',
    bg='FFEBEE'
)

# ══════════════════════════════════════════════════════════════════════════════
# Section 7: Sexual Health
# ══════════════════════════════════════════════════════════════════════════════
add_h1('Section 7: Sexual Health Impact \u2014 ICIQ-FLUTSsex (8 questions)')
add_p('All references use "sexual activity" (not "intercourse"). fs4a >= 2 = Dyspareunia flag.', space_after=6)
add_bp("Scoring: ", "FLUTSsex total 0-12. Integrated into adjunct logic (Vaginal Dilator Therapy trigger).")

# ══════════════════════════════════════════════════════════════════════════════
# Section 8: Lifestyle, Diet & Goals
# ══════════════════════════════════════════════════════════════════════════════
add_h1('Section 8: Lifestyle, Diet & Goals')

for item in [
    ("bowel_constipation: ", "Constipation frequency. >= 2 = Bowel Management Program adjunct"),
    ("symptom_triggers: ", "Activity triggers (multi-select). Includes sitting_long for pudendal neuralgia detection"),
    ("med_modify: ", "Are you changing prescribed medication due to urinary symptoms? Yes \u2192 Medication Review Referral adjunct with DIAPPERS framework guidance"),
    ("caffeine: ", "Daily caffeine intake. >= 2 = caffeine reduction recommendation"),
    ("alcohol: ", "Alcohol frequency. >= 2 = reduction advisory"),
    ("water_intake: ", "Daily water intake. Low = hydration advisory"),
    ("diet_irritants: ", "Bladder irritants (multi-select). Any = elimination trial recommendation"),
    ("prior_treatment: ", "Previous treatments (multi-select). Documented in encounter note"),
    ("cue_preference: ", 'Preferred instruction style: Biologic ("squeeze like stopping urine flow"), Imaginative ("lift a blueberry"), Breath-based ("exhale and lift"), Simple contract ("squeeze and lift"), or Default. Personalizes ALL exercise instructions.'),
    ("patient_goal: ", "Free text goal in patient's own words. Displayed in care plan header and encounter note."),
]:
    add_bp(item[0], item[1])

add_h2('Verbal Cue Personalization System')
add_p('Based on Crane, Dugan et al. (2025). Research identified 4 cueing themes with significant effectiveness differences (P < .001). The top 3 cues were "squeeze and lift" (command), "imagine stopping urine flow" (biologic function), and "breathe out, draw in and lift" (breathing).', space_after=6)
add_callout(
    'Patient selects their preferred cue style during intake. All exercise instructions throughout the care plan '
    'dynamically use their chosen language. If biologic cue is selected, a safety warning is appended: '
    '"Never practice this during actual urination."',
)

# ══════════════════════════════════════════════════════════════════════════════
# Section 9: Real-Time Consistency Checker
# ══════════════════════════════════════════════════════════════════════════════
add_h1('Section 9: Real-Time Consistency Checker (9 rules)')
add_p('Cross-instrument contradiction detection runs during intake. Inline warnings prompt the patient to review before proceeding.', space_after=6)
add_p('Rules check for contradictions across ICIQ, FLUTS, GUPI, and pain instruments. Examples: reporting "never leaks" on ICIQ but selecting leaking triggers; reporting no pain but scoring high on GUPI pain subscales; ICIQ mutual exclusion ("Never leaks" auto-deselects other options).', space_after=6)

# ══════════════════════════════════════════════════════════════════════════════
# Exercise Recommendation Logic
# ══════════════════════════════════════════════════════════════════════════════
add_h1('Exercise Recommendation Logic')
add_p('3-tier system based on Klovning severity bands. All cueing personalized per Crane/Dugan (2025). Exercise names updated for clinical context.', space_after=6)

add_h2('Tier 1: Beginner \u2014 ICIQ 13-21 (Severe/Very Severe)')
add_p('12-week program. Supine/gravity-neutral start. PT review every 3 days.', space_after=4)
for ex in [
    "1. Supine PF Activation [Start Here] \u2014 2/8/3-5s/Daily",
    "2. Diaphragmatic Breathing + PF [Foundation] \u2014 2/5/Full cycle/2x daily",
    "3. Bridge + PF [Strength] \u2014 2/10/5s/Daily",
    "4. Quick-Flick Kegels [Leak Prevention] \u2014 2/8/1s on-1s off/Daily",
    "5. Endurance Kegels [Endurance] \u2014 2/8/5s/Daily",
]:
    add_p(ex, bold=True, size=11)

add_h2('Tier 2: Moderate \u2014 ICIQ 6-12')
add_p('8-week program. Standard volume. Weekly check-in.', space_after=4)
for ex in [
    "1. Quick-Flick Kegels \u2014 3/10/1s on-1s off/Daily",
    "2. Endurance Kegels \u2014 3/10/5-10s/Daily",
    "3. Bridge + PF \u2014 2/10/5-10s/3x week",
    "4. Diaphragmatic Breathing + PF \u2014 1/5/Full cycle/Daily",
]:
    add_p(ex, bold=True, size=11)

add_h2('Tier 3: Advanced \u2014 ICIQ 1-5 (Slight)')
add_p('6-week program. Higher volume, functional focus. Bi-weekly check-in.', space_after=4)
for ex in [
    "1. Quick-Flick Kegels \u2014 3/12/1s on-1s off/Daily",
    "2. Endurance Kegels \u2014 3/12/8-10s/Daily",
    "3. Bridge + PF \u2014 3/12/10s/4x week",
    "4. Diaphragmatic Breathing \u2014 1/5/Full cycle/Daily",
]:
    add_p(ex, bold=True, size=11)

# Bridge cueing
add_h2('Bridge + PF: Detailed Cueing (per Dr. Dugan)')
add_callout(
    'Setup: Lie on back, knees bent, feet flat hip-width apart. Spine neutral \u2014 don\'t arch. '
    'Focus on core unit: PF + transverse abdominals + diaphragm + deep back muscles.\n\n'
    'Sequence: Inhale into diaphragm to prepare \u2192 Exhale + engage PF FIRST \u2192 Then lift hips \u2192 '
    'Hold 5s at top, breathe normally \u2192 Lower slowly \u2192 Fully release PF.\n\n'
    'Safety: If back hurts, core is tiring \u2014 stop immediately. Can gently tuck pelvis to initiate from hips. '
    'PF engagement matters more than height. Build to 10s hold. 10 reps \u00d7 2 sets.',
    bg='F5F5F5',
    bold_prefix='Bridge + PF Cueing: '
)

# ══════════════════════════════════════════════════════════════════════════════
# Prenatal Exercise Substitutions
# ══════════════════════════════════════════════════════════════════════════════
add_h1('Prenatal Exercise Substitution Logic')
add_p('When PRENATAL_FLAG is active, the AI plan generator automatically substitutes supine (lying-flat) exercises with incline or side-lying variants to avoid vena cava compression.', space_after=6)

add_table(
    ['Standard Exercise', 'Prenatal Substitute', 'Modification'],
    [
        ['Supine PF Activation', 'Incline PF Activation', '30-45\u00b0 incline with pillows under upper back'],
        ['Bridge + PF', 'Incline Bridge + PF', 'Incline position with towel roll under right hip'],
        ['Diaphragmatic Breathing + PF', 'Diaphragmatic Breathing + PF (Prenatal)', 'Incline or seated position'],
        ['Diaphragmatic Breathing', 'Diaphragmatic Breathing + PF (Prenatal)', 'Incline or seated position'],
    ]
)
doc.add_paragraph('')
add_p('Quick-Flick Kegels, Endurance Kegels, and Pain De-Sensitization Breathing are position-neutral and remain unchanged.', space_after=6)
add_callout(
    'Each modified exercise displays a "Prenatal-adapted" badge and inline explanation: '
    '"Modified for pregnancy \u2014 lying flat on your back after the first trimester can compress a major blood vessel. '
    'This exercise has been adjusted to a supported incline or side-lying position." '
    'A precaution block is added to the care plan. The plan is labeled "Prenatal Pelvic Floor Protocol." '
    'The PT receives a clinical alert to verify trimester appropriateness.',
    bg='E8F5E9'
)

# ══════════════════════════════════════════════════════════════════════════════
# Conditional Additions
# ══════════════════════════════════════════════════════════════════════════════
add_h1('Conditional Additions (all tiers)')
add_table(
    ['Condition', 'Action'],
    [
        ['Pain Composite >= 4', 'Add Pain De-Sensitization Breathing as FIRST exercise'],
        ['Pain >= 5 OR Dyspareunia', 'Add Vaginal Dilator Therapy adjunct'],
        ['Urgency or Mixed subtype', 'Add Bladder Retraining adjunct with full STOP technique (The Knack)'],
        ['Constipation >= 2', 'Add Bowel Management Program (footstool, fiber, water, positioning)'],
        ['Med modification = Yes', 'Add Medication Review Referral adjunct with DIAPPERS framework note'],
        ['Always', 'Add Lifestyle Adjustments (caffeine, hydration, bladder irritant guidance)'],
        ['PRENATAL_FLAG active', 'Apply prenatal exercise substitutions (see above). Add vena cava precaution.'],
    ]
)
doc.add_paragraph('')

# ══════════════════════════════════════════════════════════════════════════════
# ICD-10 Diagnosis Generation
# ══════════════════════════════════════════════════════════════════════════════
add_h1('ICD-10 Diagnosis Generation')
add_p('The AI generates up to 8 ICD-10 codes based on validated score thresholds. All diagnoses are deterministic \u2014 no generative AI in the diagnostic pipeline.', space_after=6)
add_table(
    ['Condition', 'ICD-10', 'Trigger'],
    [
        ['Stress urinary incontinence', 'N39.3', 'Subtype = Stress'],
        ['Urge urinary incontinence', 'N39.41', 'Subtype = Urge'],
        ['Mixed urinary incontinence', 'N39.46', 'Subtype = Mixed'],
        ['Continuous urinary incontinence', 'N39.45', 'Subtype = Continuous'],
        ['Pelvic pain (female)', 'R10.2', 'Pain composite >= 3'],
        ['Dyspareunia', 'N94.1', 'FLUTSsex dyspareunia flag'],
        ['Nocturia', 'R35.1', 'fl2a >= 2'],
        ['Pudendal neuralgia', 'G57.91', 'sitting_long trigger + pain >6/10 (mandatory PT review)'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# Clinical Escalation
# ══════════════════════════════════════════════════════════════════════════════
add_h1('Clinical Escalation & Risk Stratification')
add_p('Replaces DTC upsell logic. The clinical platform uses tiered escalation based on validated severity scores and behavioral indicators.', space_after=6)

add_table(
    ['Condition', 'Escalation Level', 'Action'],
    [
        ['ICIQ >= 19 OR (>= 3 avoided + elevated pain)', 'Enhanced Monitoring', 'Case flagged for enhanced monitoring. PT review frequency increased. Functional goal-setting required.'],
        ['ICIQ >= 13', 'Standard Monitoring', 'Standard PT review cadence. Outcomes tracked at 4-week and 8-week marks.'],
        ['Pudendal neuralgia indicators', 'Mandatory PT Review', 'Never auto-approved. PT must evaluate nerve involvement before progressing.'],
        ['PHQ-2 >= 3', 'Depression Risk', 'PT verifies capacity to consent. Surfaced in OAIP Red Flags dashboard.'],
        ['PRENATAL_FLAG active', 'Prenatal Protocol', 'PT reviews for trimester appropriateness. OAIP flag at MODERATE severity.'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# Encounter Note Generation
# ══════════════════════════════════════════════════════════════════════════════
add_h1('Encounter Note Generation')
add_p('The system auto-generates a CMS-compliant SOAP-format encounter note. The PT reviews, edits, and approves before it is transmitted to the referring physician via HIPAA-compliant fax.', space_after=6)

add_bp("Subjective: ", "Patient-reported symptoms, goals, pregnancy status, and relevant history.")
add_bp("Objective: ", "Validated instrument scores (ICIQ, FLUTS, GUPI, pain), incontinence subtype, severity tier, PHQ-2 result, prenatal protocol status (if applicable).")
add_bp("Assessment: ", "ICD-10 diagnoses, risk stratification (green/yellow), clinical flags.")
add_bp("Plan: ", "Exercise prescription with sets/reps/hold/frequency, adjunct recommendations, precautions, progression criteria, follow-up schedule.")
doc.add_paragraph('')
add_p('PT shorthand expansion: 12 abbreviation-to-term mappings (e.g., "PFM" \u2192 "pelvic floor muscle", "SUI" \u2192 "stress urinary incontinence") for efficient, unambiguous clinical documentation.', space_after=6)

# ══════════════════════════════════════════════════════════════════════════════
# Audit Logging
# ══════════════════════════════════════════════════════════════════════════════
add_h1('Audit Logging & OAIP Compliance')
add_p('The platform logs 24 distinct event types in an immutable audit trail. Every action taken by the patient, AI, and PT is recorded with ISO 8601 timestamps and actor identification.', space_after=6)

add_h2('OAIP Red Flags Dashboard (11 flag types)')
add_table(
    ['Flag', 'Severity', 'Trigger'],
    [
        ['SAFETY_ANSWER_CHANGED', 'CRITICAL', 'Patient changed a positive safety answer'],
        ['CLINICAL_REGRESSION_FLAG (>=7)', 'CRITICAL', 'Patient reports worsening symptoms severity >= 7'],
        ['adverse_event_report', 'HIGH', 'Patient reports an adverse event'],
        ['DEPRESSION_RISK', 'HIGH', 'PHQ-2 score >= 5'],
        ['DEPRESSION_RISK', 'MODERATE', 'PHQ-2 score 3-4'],
        ['PRENATAL_PROTOCOL_APPLIED', 'MODERATE', 'Prenatal exercise substitutions applied'],
        ['DEFAULT_FAX_APPLIED', 'LOW', 'Default fax used (NPI had no fax on file)'],
        ['EXERCISE_PAIN_REPORT', 'Varies', 'Patient reports pain during exercises'],
        ['CLINICAL_REGRESSION_FLAG (<7)', 'MODERATE', 'Patient reports worsening symptoms severity < 7'],
        ['clinical_review_request', 'HIGH', 'Patient requests AI opt-out'],
        ['CARE_PLAN_DOWNLOADED', 'LOW', 'Patient exported care plan for EHR sharing'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# HIPAA Reminder
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph('')
add_h1('HIPAA PHI Handling \u2014 Reviewer Responsibilities')
add_callout(
    'All data in the PT review interface is Protected Health Information (PHI). The audit log captures PHI '
    '(patient names, provider names, fax numbers, clinical details) as part of the treatment record. '
    'When using Auditor Mode, 17 PHI-sensitive field types are automatically masked with hashed identifiers. '
    'This provides display-level de-identification for regulatory review but is not HIPAA Safe Harbor de-identification. '
    'Do not copy, screenshot, or transmit PHI outside the platform. '
    'All reviewer actions are logged and available for OAIP audit.',
    bg='FFEBEE'
)

# Footer
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('CONFIDENTIAL \u2014 Utah OAIP Regulatory Sandbox Pilot')
run.font.size = Pt(9); run.font.color.rgb = GRAY; run.italic = True
p2 = doc.add_paragraph()
run2 = p2.add_run('For clinical review by Dr. Nicole Dugan and authorized Expect clinical team members only.')
run2.font.size = Pt(9); run2.font.color.rgb = GRAY; run2.italic = True

output_path = '/Users/Dara Cook/Downloads/expect-utah/utah-pt-logic-review.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
